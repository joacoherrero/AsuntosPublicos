import os
import json
from datetime import datetime, timedelta
import pdfplumber
import re
import pandas as pd
from typing import Dict, List, Optional, Any, Tuple
import logging
import csv
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import tempfile
from docx import Document

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('boletin_processor.log'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

class BoletinDownloader:
    """Clase para descargar el Boletín Oficial"""
    
    def __init__(self):
        self.base_url = "https://www.boletinoficial.gob.ar"
        self.pdf_url_template = "https://www.boletinoficial.gob.ar/pdf/pdfPorNombre/{fecha}"
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
    
    def obtener_ultimo_boletin(self) -> Tuple[str, str]:
        """
        Obtiene la URL del último Boletín Oficial publicado
        
        Returns:
            Tuple[str, str]: (URL del PDF, nombre del archivo)
        """
        try:
            logger.info("Obteniendo URL del último Boletín Oficial...")
            
            # Obtener la fecha actual
            fecha_actual = datetime.now()
            
            # Si es fin de semana, retroceder al viernes
            while fecha_actual.weekday() > 4:  # 5 = Sábado, 6 = Domingo
                fecha_actual = fecha_actual - timedelta(days=1)
            
            # Formatear la fecha para la URL
            fecha_str = fecha_actual.strftime("%Y%m%d")
            
            # Construir la URL del PDF
            pdf_url = self.pdf_url_template.format(fecha=fecha_str)
            nombre_archivo = f"boletin_{fecha_str}.pdf"
            
            # Verificar si el PDF existe
            response = requests.head(pdf_url, headers=self.headers)
            if response.status_code == 404:
                # Si no existe, probar con el día anterior
                fecha_actual = fecha_actual - timedelta(days=1)
                fecha_str = fecha_actual.strftime("%Y%m%d")
                pdf_url = self.pdf_url_template.format(fecha=fecha_str)
                nombre_archivo = f"boletin_{fecha_str}.pdf"
                
                # Verificar nuevamente
                response = requests.head(pdf_url, headers=self.headers)
                if response.status_code == 404:
                    raise ValueError(f"No se encontró el Boletín para la fecha {fecha_str}")
            
            logger.info(f"URL del Boletín encontrada: {pdf_url}")
            return pdf_url, nombre_archivo
            
        except Exception as e:
            logger.error(f"Error al obtener la URL del Boletín: {str(e)}")
            raise
    
    def descargar_boletin(self, url: str, nombre_archivo: str) -> str:
        """
        Descarga el PDF del Boletín Oficial
        
        Args:
            url (str): URL del PDF
            nombre_archivo (str): Nombre para guardar el archivo
            
        Returns:
            str: Ruta al archivo descargado
        """
        try:
            logger.info(f"Descargando Boletín desde {url}...")
            
            # Crear directorio temporal si no existe
            temp_dir = "temp_boletines"
            os.makedirs(temp_dir, exist_ok=True)
            
            # Ruta completa para el archivo
            ruta_archivo = os.path.join(temp_dir, nombre_archivo)
            
            # Descargar el archivo
            response = requests.get(url, headers=self.headers, stream=True)
            response.raise_for_status()
            
            # Guardar el archivo
            with open(ruta_archivo, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            logger.info(f"Boletín descargado exitosamente: {ruta_archivo}")
            return ruta_archivo
            
        except Exception as e:
            logger.error(f"Error al descargar el Boletín: {str(e)}")
            raise

    def descargar_primera_seccion(self, fecha: str = None) -> str:
        """
        Descarga específicamente la primera sección del Boletín Oficial para una fecha dada
        
        Args:
            fecha (str, optional): Fecha en formato YYYYMMDD. Si no se proporciona, se usa la fecha actual
            
        Returns:
            str: Ruta al archivo descargado
        """
        try:
            if fecha is None:
                fecha_actual = datetime.now()
                # Si es fin de semana, retroceder al viernes
                while fecha_actual.weekday() > 4:  # 5 = Sábado, 6 = Domingo
                    fecha_actual = fecha_actual - timedelta(days=1)
                fecha = fecha_actual.strftime("%Y%m%d")

            logger.info(f"Descargando primera sección del Boletín para la fecha {fecha}...")
            
            # Headers específicos para la descarga
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'es-ES,es;q=0.8,en-US;q=0.5,en;q=0.3',
                'Accept-Encoding': 'gzip, deflate, br',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1'
            }
            
            # Iniciar sesión y obtener cookies
            session = requests.Session()
            
            # 1. Acceder a la página principal para obtener cookies iniciales
            logger.info("Accediendo a la página principal...")
            response = session.get(self.base_url, headers=headers)
            response.raise_for_status()
            
            # 2. Acceder a la sección primera
            main_url = f"{self.base_url}/seccion/primera"
            logger.info(f"Accediendo a la sección primera: {main_url}")
            response = session.get(main_url, headers=headers)
            response.raise_for_status()
            
            # Analizar la página HTML
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Buscar el botón de descarga
            download_button = soup.find('button', {
                'class': 'btn-circle',
                'onclick': lambda x: 'descargarPDFSeccion' in str(x) if x else False
            })
            
            if not download_button:
                logger.error("No se encontró el botón de descarga en la página")
                raise ValueError("No se encontró el botón de descarga en la página")
            
            # Extraer los parámetros de la función onclick
            onclick = download_button.get('onclick', '')
            logger.info(f"Función onclick encontrada: {onclick}")
            
            # La función tiene el formato: descargarPDFSeccion("primera","20250523", "/pdf/download_section")
            import re
            params = re.findall(r'"([^"]*)"', onclick)
            if len(params) < 3:
                raise ValueError(f"No se pudieron extraer los parámetros de descarga: {onclick}")
            
            seccion, fecha_descarga, endpoint = params[:3]
            
            # Construir la URL de descarga
            download_url = f"{self.base_url}/seccion/primera/pdf/{fecha}"
            nombre_archivo = f"boletin_primera_seccion_{fecha}.pdf"
            
            logger.info(f"URL de descarga construida: {download_url}")
            
            # Configurar headers específicos para la descarga del PDF
            headers.update({
                'Accept': 'application/pdf',
                'Content-Type': 'application/pdf'
            })
            
            # Intentar la descarga
            logger.info("Iniciando descarga del PDF...")
            response = session.get(download_url, headers=headers, stream=True)
            response.raise_for_status()
            
            # Verificar que la respuesta es un PDF
            content_type = response.headers.get('content-type', '')
            if 'application/pdf' not in content_type.lower():
                logger.error(f"La respuesta no es un PDF. Content-Type: {content_type}")
                logger.error(f"Primeros 1000 bytes de la respuesta: {response.content[:1000]}")
                raise ValueError(f"La respuesta no es un PDF (Content-Type: {content_type})")
            
            # Crear directorio temporal si no existe
            temp_dir = "temp_boletines"
            os.makedirs(temp_dir, exist_ok=True)
            
            # Ruta completa para el archivo
            ruta_archivo = os.path.join(temp_dir, nombre_archivo)
            
            # Guardar el archivo
            with open(ruta_archivo, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            logger.info(f"Boletín descargado exitosamente: {ruta_archivo}")
            return ruta_archivo
            
        except Exception as e:
            logger.error(f"Error al descargar la primera sección del Boletín: {str(e)}")
            raise

    def descargar_primera_seccion_selenium(self, fecha: str = None) -> str:
        """
        Descarga la primera sección del Boletín Oficial usando Selenium
        
        Args:
            fecha (str, optional): Fecha en formato YYYYMMDD. Si no se proporciona, se usa la fecha actual
            
        Returns:
            str: Ruta al archivo descargado
        """
        try:
            import undetected_chromedriver as uc
            from selenium.webdriver.common.by import By
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            import time
            
            if fecha is None:
                fecha_actual = datetime.now()
                # Si es fin de semana, retroceder al viernes
                while fecha_actual.weekday() > 4:  # 5 = Sábado, 6 = Domingo
                    fecha_actual = fecha_actual - timedelta(days=1)
                fecha = fecha_actual.strftime("%Y%m%d")

            logger.info(f"Descargando primera sección del Boletín para la fecha {fecha} usando Selenium...")
            
            # Configurar opciones de Chrome
            chrome_options = uc.ChromeOptions()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            
            # Configurar preferencias de descarga
            temp_dir = os.path.abspath("temp_boletines")
            os.makedirs(temp_dir, exist_ok=True)
            
            chrome_options.add_experimental_option('prefs', {
                'download.default_directory': temp_dir,
                'download.prompt_for_download': False,
                'download.directory_upgrade': True,
                'safebrowsing.enabled': True,
                'plugins.always_open_pdf_externally': True  # Descargar PDFs en lugar de abrirlos
            })
            
            # Iniciar el navegador
            driver = uc.Chrome(options=chrome_options)
            wait = WebDriverWait(driver, 10)
            
            try:
                # Acceder a la página principal
                logger.info("Accediendo a la página principal...")
                driver.get(f"{self.base_url}/seccion/primera")
                
                # Esperar a que el botón de descarga esté presente y sea clickeable
                download_button = wait.until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, 'button.btn-circle[onclick*="descargarPDFSeccion"]'))
                )
                
                # Obtener el valor onclick del botón
                onclick_value = download_button.get_attribute('onclick')
                logger.info(f"Valor onclick del botón: {onclick_value}")
                
                # Extraer la URL de descarga del valor onclick
                import re
                match = re.search(r"window\.open\('([^']+)'", onclick_value)
                if not match:
                    raise ValueError("No se pudo encontrar la URL de descarga en el botón")
                
                download_url = match.group(1)
                if not download_url.startswith('http'):
                    download_url = f"{self.base_url}{download_url}"
                
                logger.info(f"URL de descarga: {download_url}")
                
                # Navegar directamente a la URL de descarga
                driver.get(download_url)
                
                # Esperar a que se complete la descarga
                nombre_archivo = f"boletin_primera_seccion_{fecha}.pdf"
                ruta_archivo = os.path.join(temp_dir, nombre_archivo)
                
                # Esperar hasta 30 segundos por la descarga
                timeout = 30
                while timeout > 0:
                    if os.path.exists(ruta_archivo) and os.path.getsize(ruta_archivo) > 0:
                        # Verificar que sea un PDF válido
                        try:
                            with open(ruta_archivo, 'rb') as f:
                                header = f.read(4)
                                if header.startswith(b'%PDF'):
                                    logger.info(f"Boletín descargado exitosamente: {ruta_archivo}")
                                    return ruta_archivo
                                else:
                                    logger.warning("El archivo descargado no es un PDF válido")
                        except Exception as e:
                            logger.warning(f"Error al verificar el archivo: {str(e)}")
                    time.sleep(1)
                    timeout -= 1
                
                raise TimeoutError("Se agotó el tiempo de espera para la descarga")
                
            finally:
                driver.quit()
                logger.info("Navegador cerrado")
            
        except Exception as e:
            logger.error(f"Error al descargar la primera sección del Boletín usando Selenium: {str(e)}")
            raise

class BoletinPDFProcessor:
    def __init__(self):
        """Inicializa el procesador de PDF del Boletín Oficial"""
        self.output_dir = "datos_boletin"
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Cargar diccionario y cuentas
        self.temas = self._cargar_diccionario()
        self.cuentas = self._cargar_cuentas()
        
        # Patrones de expresiones regulares para el boletín
        self.patrones = {
            'numero_boletin': r'Boletín\s+Oficial\s+N[°º]\s+(\d+\.?\d*)',
            'fecha_boletin': r'(?:Lunes|Martes|Miércoles|Jueves|Viernes)\s+(\d{1,2})\s+de\s+([A-Za-zÁÉÍÓÚáéíóúñÑ]+)\s+de\s+(\d{4})',
            'codigo_hash': r'#[IF]\d+[IF]#',
            'codigo_publicacion': r'e\.\s+\d{2}/\d{2}/\d{4}\s+N°\s+\d+/\d{2}\s+v\.\s+\d{2}/\d{2}/\d{4}',
            'expediente': r'EX-\d{4}-\d+-[A-Z]+-[A-Z]+#[A-Z]+',
            'identificador': r'[A-Z]+-\d{4}-\d+-[A-Z]+-[A-Z]+',
        }
        
        # Tipos de documentos soportados
        self.tipos_documentos = {
            'LEY': r'LEY\s+(?:N[°º]\s*)?(\d+(?:/\d{4})?)',
            'DECRETO': r'DECRETO\s+(?:N[°º]\s*)?(\d+(?:/\d{4})?)',
            'RESOLUCIÓN': r'RESOLUCIÓN\s+(?:N[°º]\s*)?(\d+(?:/\d{4})?)',
            'DISPOSICIÓN': r'DISPOSICIÓN\s+(?:N[°º]\s*)?(\d+(?:/\d{4})?)',
            'DECISIÓN ADMINISTRATIVA': r'DECISIÓN\s+ADMINISTRATIVA\s+(?:N[°º]\s*)?(\d+(?:/\d{4})?)'
        }
    
    def _cargar_diccionario(self) -> Dict:
        """Carga los temas y palabras clave del diccionario Excel"""
        try:
            df = pd.read_excel('Diccionario.xlsx', engine='openpyxl')
            temas = {}
            
            # Iterar sobre las filas del DataFrame
            for _, row in df.iterrows():
                if pd.notna(row.iloc[0]):  # Si hay tema
                    tema = str(row.iloc[0]).strip()
                    if tema not in temas:
                        temas[tema] = {
                            'tema': tema,
                            'palabras_clave': []
                        }
                    
                    # Si hay palabra clave, agregarla
                    if len(row) > 1 and pd.notna(row.iloc[1]):
                        palabra_clave = str(row.iloc[1]).strip().lower()
                        temas[tema]['palabras_clave'].append(palabra_clave)
            
            logger.info(f"Se cargaron {len(temas)} temas del diccionario")
            return temas
        except Exception as e:
            logger.error(f"Error al cargar el diccionario: {str(e)}")
            return {}

    def _cargar_cuentas(self) -> Dict:
        """Carga la información de cuentas y sus temas asociados desde el Excel"""
        try:
            df = pd.read_excel('Cuentas.xlsx', engine='openpyxl')
            cuentas = {}
            
            # Procesar cada fila (empresa)
            for idx, row in df.iterrows():
                try:
                    empresa = row['Empresa']
                    if pd.isna(empresa):
                        continue
                    
                    empresa = str(empresa).strip()
                    if empresa not in cuentas:
                        cuentas[empresa] = {
                            'nombre': empresa,
                            'temas': []
                        }
                    
                    # Procesar cada columna de temáticas
                    for columna in df.columns:
                        if columna != 'Empresa' and pd.notna(row[columna]):
                            tema = str(row[columna]).strip()
                            if tema and tema not in cuentas[empresa]['temas']:
                                cuentas[empresa]['temas'].append(tema)
                
                except Exception as row_error:
                    logger.error(f"Error procesando fila {idx}: {str(row_error)}")
                    continue
            
            logger.info(f"Se cargaron {len(cuentas)} cuentas")
            return cuentas
        except Exception as e:
            logger.error(f"Error al cargar el archivo de cuentas: {str(e)}")
            return {}

    def _clasificar_documento(self, texto: str) -> List[Dict]:
        """Clasifica un documento según los temas del diccionario"""
        texto_lower = texto.lower()
        temas_encontrados = []
        
        for tema_data in self.temas.values():
            # Por cada palabra clave del tema, verificar si está contenida en el texto
            for palabra_clave in tema_data['palabras_clave']:
                if palabra_clave in texto_lower:
                    tema_encontrado = {
                        'tema': tema_data['tema'],
                        'palabra_encontrada': palabra_clave
                    }
                    temas_encontrados.append(tema_encontrado)
                    break
        
        return temas_encontrados

    def _encontrar_cuentas_interesadas(self, temas_documento: List[Dict]) -> List[str]:
        """Encuentra las cuentas interesadas en los temas del documento"""
        cuentas_interesadas = []
        
        for cuenta, data in self.cuentas.items():
            for tema_doc in temas_documento:
                if tema_doc['tema'] in data['temas']:
                    cuentas_interesadas.append(cuenta)
                    break
        
        return list(set(cuentas_interesadas))  # Eliminar duplicados

    def procesar_pdf(self, pdf_path: str) -> Optional[Dict]:
        """
        Procesa el PDF del Boletín Oficial y extrae su contenido
        
        Args:
            pdf_path (str): Ruta al archivo PDF
            
        Returns:
            Optional[Dict]: Diccionario con el contenido procesado o None si hay error
        """
        pdf = None
        try:
            logger.info(f"Procesando PDF: {pdf_path}")
            
            if not os.path.exists(pdf_path):
                logger.error(f"No se encontró el archivo: {pdf_path}")
                return None
            
            metadata = {
                "fecha_proceso": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "archivo": pdf_path,
                "numero_boletin": None,
                "fecha_boletin": None,
                "documentos": []
            }
            
            pdf = pdfplumber.open(pdf_path)
            logger.info(f"PDF abierto correctamente - {len(pdf.pages)} páginas")
            
            # Procesar sumario (primeras páginas)
            sumario = self._extraer_sumario(pdf.pages[:3])
            metadata["sumario"] = sumario
            
            # Procesar cada página
            documento_actual = None
            texto_documento = []
            
            for i, pagina in enumerate(pdf.pages):
                logger.info(f"Procesando página {i+1}/{len(pdf.pages)}...")
                texto = pagina.extract_text()
                
                # Extraer metadata del boletín (solo en primera página)
                if i == 0:
                    metadata.update(self._extraer_metadata_boletin(texto))
                
                # Procesar el contenido de la página
                for linea in texto.split('\n'):
                    nuevo_documento = self._detectar_inicio_documento(linea)
                    
                    if nuevo_documento:
                        # Guardar documento anterior si existe
                        if documento_actual and texto_documento:
                            doc_procesado = self._procesar_documento(
                                documento_actual,
                                '\n'.join(texto_documento)
                            )
                            metadata["documentos"].append(doc_procesado)
                        
                        documento_actual = nuevo_documento
                        texto_documento = [linea]
                    else:
                        if documento_actual:
                            texto_documento.append(linea)
            
            # Procesar último documento
            if documento_actual and texto_documento:
                doc_procesado = self._procesar_documento(
                    documento_actual,
                    '\n'.join(texto_documento)
                )
                metadata["documentos"].append(doc_procesado)
            
            # Guardar resultados
            self._guardar_resultados(metadata)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error al procesar el PDF: {str(e)}", exc_info=True)
            return None
        finally:
            if pdf:
                logger.info("Cerrando archivo PDF...")
                pdf.close()
                logger.info("PDF cerrado correctamente")

    def _extraer_sumario(self, paginas_iniciales: List) -> List[Dict]:
        """Extrae el sumario de las primeras páginas"""
        sumario = []
        en_sumario = False
        
        for pagina in paginas_iniciales:
            texto = pagina.extract_text()
            
            # Detectar inicio del sumario
            if "SUMARIO" in texto and not en_sumario:
                en_sumario = True
                continue
            
            if not en_sumario:
                continue
            
            # Procesar líneas del sumario
            for linea in texto.split('\n'):
                if not linea.strip():
                    continue
                    
                # Detectar entradas del sumario (ejemplo: "Decreto 350/2025...")
                for tipo in self.tipos_documentos.keys():
                    if tipo in linea:
                        match = re.search(fr'{tipo}\s+(\d+/\d{4}).*?\.+\s*pág\.\s*(\d+)', linea)
                        if match:
                            sumario.append({
                                "tipo": tipo,
                                "numero": match.group(1),
                                "pagina": int(match.group(2)),
                                "referencia": linea.strip()
                            })
                            break
                
                # Detectar fin del sumario
                if "Primera Sección" in linea:
                    return sumario
        
        return sumario
    
    def _extraer_metadata_boletin(self, texto: str) -> Dict:
        """Extrae la metadata general del boletín"""
        metadata = {}
        
        # Extraer número de boletín
        match = re.search(self.patrones['numero_boletin'], texto)
        if match:
            metadata['numero_boletin'] = match.group(1)
        
        # Extraer fecha
        match = re.search(self.patrones['fecha_boletin'], texto)
        if match:
            metadata['fecha_boletin'] = f"{match.group(1)} de {match.group(2)} de {match.group(3)}"
        
        return metadata
    
    def _detectar_inicio_documento(self, linea: str) -> Optional[str]:
        """Detecta si una línea corresponde al inicio de un nuevo documento"""
        for tipo, patron in self.tipos_documentos.items():
            if re.match(patron, linea, re.IGNORECASE):
                return tipo
        return None
    
    def _procesar_documento(self, tipo: str, texto: str) -> Dict[str, Any]:
        """Procesa un documento individual y extrae su información"""
        documento = {
            "tipo_documento": tipo,
            "numero_documento": self._extraer_numero_documento(tipo, texto),
            "identificador": self._extraer_identificador(texto),
            "fecha": self._extraer_fecha(texto),
            "titulo": self._extraer_titulo(texto),
            "organismo_emisor": self._extraer_organismo(texto),
            "contenido": texto,
            "firmantes": self._extraer_firmantes(texto),
            "codigo_publicacion": self._extraer_codigo_publicacion(texto),
            "codigo_hash": self._extraer_codigo_hash(texto),
            "tiene_anexo_web": "ANEXO" in texto and "web" in texto.lower()
        }
        
        return documento
    
    def _extraer_numero_documento(self, tipo: str, texto: str) -> Optional[str]:
        """Extrae el número del documento"""
        patron = self.tipos_documentos.get(tipo)
        if patron:
            match = re.search(patron, texto)
            if match:
                return match.group(1)
        return None
    
    def _extraer_identificador(self, texto: str) -> Optional[str]:
        """Extrae el identificador normativo"""
        match = re.search(self.patrones['identificador'], texto)
        return match.group() if match else None
    
    def _extraer_fecha(self, texto: str) -> Optional[str]:
        """Extrae la fecha del documento"""
        patrones = [
            r'Ciudad de Buenos Aires,\s*(\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de\s+\d{4})',
            r'(\d{2}/\d{2}/\d{4})'
        ]
        
        for patron in patrones:
            match = re.search(patron, texto)
            if match:
                return match.group(1)
        return None
    
    def _extraer_titulo(self, texto: str) -> Optional[str]:
        """Extrae el título o descripción del documento"""
        lineas = texto.split('\n')
        for i, linea in enumerate(lineas[:5]):  # Buscar en las primeras 5 líneas
            if '-' in linea and len(linea) > 10:
                return linea.strip()
        return None
    
    def _extraer_organismo(self, texto: str) -> Optional[str]:
        """Extrae el organismo emisor"""
        organismos = [
            r'MINISTERIO\s+DE\s+[A-ZÁÉÍÓÚÑ\s]+',
            r'SECRETARÍA\s+[A-ZÁÉÍÓÚÑ\s]+',
            r'PRESIDENCIA\s+DE\s+LA\s+NACIÓN'
        ]
        
        for patron in organismos:
            match = re.search(patron, texto)
            if match:
                return match.group().strip()
        return None
    
    def _extraer_firmantes(self, texto: str) -> List[str]:
        """Extrae los nombres de los firmantes"""
        firmantes = []
        lineas = texto.split('\n')
        
        # Buscar líneas que parezcan nombres al final del documento
        for linea in lineas[-10:]:  # Últimas 10 líneas
            if re.match(r'^[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)+$', linea):
                firmantes.append(linea.strip())
        
        return firmantes
    
    def _extraer_codigo_publicacion(self, texto: str) -> Optional[str]:
        """Extrae el código de publicación"""
        match = re.search(self.patrones['codigo_publicacion'], texto)
        return match.group() if match else None
    
    def _extraer_codigo_hash(self, texto: str) -> Optional[str]:
        """Extrae el código hash interno"""
        match = re.search(self.patrones['codigo_hash'], texto)
        return match.group() if match else None
    
    def _guardar_resultados(self, metadata: Dict) -> None:
        """Guarda los resultados en diferentes formatos"""
        try:
            # Generar nombre base con fecha y hora
            fecha_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_nombre = f"boletin_{metadata.get('numero_boletin', fecha_str)}"
            
            # Asegurarse de que el directorio existe y tenemos permisos
            os.makedirs(self.output_dir, exist_ok=True)
            
            # Preparar datos para el CSV
            documentos_data = []
            for doc in metadata['documentos']:
                # Limpiar y escapar el contenido
                contenido_limpio = doc['contenido'].replace('\r', ' ').replace('\n', '\\n')
                
                # Clasificar el documento
                temas_encontrados = self._clasificar_documento(contenido_limpio)
                cuentas_interesadas = self._encontrar_cuentas_interesadas(temas_encontrados)
                
                # Crear un diccionario con todos los campos relevantes
                documento = {
                    'tipo': doc['tipo_documento'],
                    'numero': doc['numero_documento'],
                    'fecha': doc['fecha'],
                    'titulo': doc['titulo'].replace('\r', ' ').replace('\n', ' ') if doc['titulo'] else '',
                    'organismo': doc['organismo_emisor'],
                    'identificador': doc['identificador'],
                    'codigo_publicacion': doc['codigo_publicacion'],
                    'codigo_hash': doc['codigo_hash'],
                    'firmantes': '; '.join(doc['firmantes']) if doc['firmantes'] else '',
                    'tiene_anexo_web': doc['tiene_anexo_web'],
                    'temas_detectados': '; '.join(t['tema'] for t in temas_encontrados),
                    'palabras_clave': '; '.join(t['palabra_encontrada'] for t in temas_encontrados),
                    'cuentas_interesadas': '; '.join(cuentas_interesadas),
                    'contenido_completo': contenido_limpio
                }
                documentos_data.append(documento)
            
            # Crear DataFrame
            df = pd.DataFrame(documentos_data)
            
            # Reordenar columnas para mejor legibilidad
            columnas_orden = [
                'tipo', 'numero', 'fecha', 'titulo', 'organismo', 
                'identificador', 'codigo_publicacion', 'codigo_hash',
                'firmantes', 'tiene_anexo_web', 'temas_detectados',
                'palabras_clave', 'cuentas_interesadas', 'contenido_completo'
            ]
            df = df[columnas_orden]
            
            # Intentar guardar en diferentes formatos
            try:
                # 1. Guardar como TSV (Tab-Separated Values)
                tsv_path = os.path.join(self.output_dir, f"{base_nombre}_documentos.tsv")
                df.to_csv(tsv_path, 
                         index=False,
                         encoding='utf-8-sig',
                         sep='\t',
                         quoting=csv.QUOTE_MINIMAL)
                logger.info(f"Datos guardados en TSV: {tsv_path}")
            except Exception as e:
                logger.error(f"Error al guardar TSV: {str(e)}")
            
            try:
                # 2. Guardar como Excel
                excel_path = os.path.join(self.output_dir, f"{base_nombre}_documentos.xlsx")
                df.to_excel(excel_path, index=False, engine='openpyxl')
                logger.info(f"Datos guardados en Excel: {excel_path}")
            except Exception as e:
                logger.error(f"Error al guardar Excel: {str(e)}")
            
            try:
                # 3. Guardar JSON (como respaldo)
                json_path = os.path.join(self.output_dir, f"{base_nombre}_metadata.json")
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, ensure_ascii=False, indent=2)
                logger.info(f"Metadata guardada en JSON: {json_path}")
            except Exception as e:
                logger.error(f"Error al guardar JSON: {str(e)}")
            
            # Generar resumen por cuenta
            self._generar_resumen_por_cuenta(df, base_nombre)
            
            # Mostrar las primeras filas como verificación
            logger.info("\nPrimeros documentos procesados:")
            logger.info(df[['tipo', 'numero', 'fecha', 'titulo', 'temas_detectados']].head().to_string())
            
        except Exception as e:
            logger.error(f"Error al guardar resultados: {str(e)}")
            raise

    def _generar_resumen_por_cuenta(self, df: pd.DataFrame, base_nombre: str) -> None:
        """Genera un resumen Excel y Word por cada cuenta con sus documentos relevantes"""
        try:
            # Crear directorio para reportes por cuenta
            dir_reportes = os.path.join(self.output_dir, f"{base_nombre}_reportes_cuenta")
            os.makedirs(dir_reportes, exist_ok=True)
            
            # Para cada cuenta
            for cuenta in self.cuentas.keys():
                try:
                    # Filtrar documentos relevantes para esta cuenta
                    docs_cuenta = df[df['cuentas_interesadas'].str.contains(cuenta, na=False)]
                    
                    if not docs_cuenta.empty:
                        # 1. Guardar reporte Excel
                        nombre_excel = f"reporte_{cuenta.replace(' ', '_').replace('/', '_')}.xlsx"
                        ruta_excel = os.path.join(dir_reportes, nombre_excel)
                        
                        try:
                            docs_cuenta.to_excel(ruta_excel, 
                                              index=False,
                                              engine='openpyxl')
                            logger.info(f"Generado reporte Excel para {cuenta} con {len(docs_cuenta)} documentos")
                        except Exception as e:
                            logger.warning(f"Error al guardar reporte Excel para {cuenta}: {str(e)}")
                        
                        # 2. Generar reporte Word
                        try:
                            doc = Document()
                            
                            # Título
                            doc.add_heading(f'Reporte de Boletín Oficial - {cuenta}', 0)
                            
                            # Información general
                            doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
                            doc.add_paragraph(f'Total de documentos relevantes: {len(docs_cuenta)}')
                            doc.add_paragraph()
                            
                            # Agregar cada documento
                            for _, documento in docs_cuenta.iterrows():
                                # Título del documento
                                doc.add_heading(f'{documento["tipo"]} {documento["numero"] or "S/N"}', level=1)
                                
                                # Metadata básica
                                p = doc.add_paragraph()
                                p.add_run('Fecha: ').bold = True
                                p.add_run(f'{documento["fecha"] or "No especificada"}')
                                
                                p = doc.add_paragraph()
                                p.add_run('Organismo: ').bold = True
                                p.add_run(f'{documento["organismo"] or "No especificado"}')
                                
                                if documento["identificador"]:
                                    p = doc.add_paragraph()
                                    p.add_run('Identificador: ').bold = True
                                    p.add_run(documento["identificador"])
                                
                                # Temas y palabras clave
                                if documento["temas_detectados"]:
                                    p = doc.add_paragraph()
                                    p.add_run('Temas detectados: ').bold = True
                                    p.add_run(documento["temas_detectados"])
                                
                                if documento["palabras_clave"]:
                                    p = doc.add_paragraph()
                                    p.add_run('Palabras clave: ').bold = True
                                    p.add_run(documento["palabras_clave"])
                                
                                # Título/descripción del documento
                                if documento["titulo"]:
                                    doc.add_heading('Descripción', level=2)
                                    doc.add_paragraph(documento["titulo"])
                                
                                # Contenido completo
                                doc.add_heading('Contenido Completo', level=2)
                                contenido = documento["contenido_completo"].replace('\\n', '\n')
                                doc.add_paragraph(contenido)
                                
                                # Separador entre documentos
                                doc.add_paragraph('_' * 50)
                                doc.add_paragraph()
                            
                            # Guardar documento Word
                            nombre_word = f"reporte_detallado_{cuenta.replace(' ', '_').replace('/', '_')}.docx"
                            ruta_word = os.path.join(dir_reportes, nombre_word)
                            doc.save(ruta_word)
                            logger.info(f"Generado reporte Word detallado para {cuenta}")
                            
                        except Exception as e:
                            logger.warning(f"Error al generar reporte Word para {cuenta}: {str(e)}")
                
                except Exception as e:
                    logger.warning(f"Error procesando cuenta {cuenta}: {str(e)}")
                    continue
        
        except Exception as e:
            logger.error(f"Error al generar resúmenes por cuenta: {str(e)}")
            # No re-raise aquí para permitir que el proceso principal continúe

def main():
    """Función principal"""
    try:
        # Ruta al archivo de ejemplo
        pdf_path = r"C:\Users\herre\OneDrive\Escritorio\Cursor\Proyecto Contenido\BO.pdf"
        
        if not os.path.exists(pdf_path):
            logger.error(f"\n❌ No se encontró el archivo: {pdf_path}")
            return
            
        # Procesar el PDF
        processor = BoletinPDFProcessor()
        metadata = processor.procesar_pdf(pdf_path)
        
        if metadata:
            logger.info(f"\n✨ Proceso completado. Se encontraron {len(metadata['documentos'])} documentos.")
            logger.info(f"Boletín N° {metadata.get('numero_boletin')} del {metadata.get('fecha_boletin')}")
        else:
            logger.error("\n❌ No se pudo procesar el PDF.")
            
    except Exception as e:
        logger.error(f"\n❌ Error en el proceso principal: {str(e)}")
        raise

if __name__ == "__main__":
    main() 