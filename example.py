# Importación de las clases necesarias del proyecto
from rss_reader import RSSReader
from collections import defaultdict
import re
import time
import pandas as pd
import concurrent.futures
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
import gspread
from google.oauth2.service_account import Credentials
import logging
import sys

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)

def cargar_diccionario():
    """Carga los temas y palabras clave del diccionario Excel"""
    try:
        df = pd.read_excel('Diccionario.xlsx', engine='openpyxl')
        temas = {}
        
        # Iterar sobre las filas del DataFrame
        for _, row in df.iterrows():
            if pd.notna(row.iloc[0]):  # Si hay tema
                tema = str(row.iloc[0]).strip()
                if tema not in temas:
                    temas[tema] = {
                        'tema': tema,
                        'palabras_clave': []
                    }
                
                # Si hay palabra clave, agregarla
                if len(row) > 1 and pd.notna(row.iloc[1]):
                    palabra_clave = str(row.iloc[1]).strip().lower()
                    temas[tema]['palabras_clave'].append(palabra_clave)
        
        return temas
    except Exception as e:
        print(f"❌ Error al cargar el diccionario: {str(e)}")
        return {}

def cargar_noticias_sheets():
    """Carga las noticias del día desde Google Sheets"""
    try:
        print("\n📊 Intentando conectar con Google Sheets...")
        
        # Verificar si existe el archivo de credenciales
        if not os.path.exists('credentials.json'):
            print("❌ No se encontró el archivo credentials.json")
            return []
        
        # Configurar credenciales
        scopes = [
            'https://www.googleapis.com/auth/spreadsheets.readonly',
            'https://www.googleapis.com/auth/drive.readonly'
        ]
        
        print("🔑 Cargando credenciales desde credentials.json...")
        try:
            credentials = Credentials.from_service_account_file('credentials.json', scopes=scopes)
            print("✅ Credenciales cargadas correctamente")
        except Exception as e:
            print(f"❌ Error al cargar credenciales: {str(e)}")
            return []
        
        print("🔄 Autorizando con Google...")
        try:
            gc = gspread.authorize(credentials)
            print("✅ Autorización exitosa")
        except Exception as e:
            print(f"❌ Error en la autorización: {str(e)}")
            return []
        
        # ID de la hoja de cálculo - REEMPLAZA ESTE ID CON EL DE TU HOJA
        SPREADSHEET_ID = '1MwmcgWqaJzTimdsjvAyLNRZTKxqfqjG9o8PygxKeLrA'
        
        print(f"📝 Accediendo a la hoja de cálculo (ID: {SPREADSHEET_ID})...")
        try:
            spreadsheet = gc.open_by_key(SPREADSHEET_ID)
            print("✅ Hoja de cálculo abierta correctamente")
        except gspread.exceptions.SpreadsheetNotFound:
            print("❌ No se encontró la hoja de cálculo con ese ID")
            return []
        except Exception as e:
            print(f"❌ Error al abrir la hoja de cálculo: {str(e)}")
            return []
        
        print("📊 Accediendo a la hoja 'Noticias'...")
        try:
            worksheet = spreadsheet.worksheet('Noticias')
            print("✅ Hoja 'Noticias' encontrada")
        except gspread.exceptions.WorksheetNotFound:
            print("❌ No se encontró la hoja 'Noticias' en el documento")
            return []
        except Exception as e:
            print(f"❌ Error al acceder a la hoja 'Noticias': {str(e)}")
            return []
        
        print("📥 Leyendo datos...")
        try:
            data = worksheet.get_all_values()
            print(f"✅ Se leyeron {len(data)} filas de datos")
        except Exception as e:
            print(f"❌ Error al leer los datos: {str(e)}")
            return []
        
        # Convertir a DataFrame
        df = pd.DataFrame(data[1:], columns=data[0])  # Asumiendo que la primera fila son los encabezados
        
        # Filtrar por la fecha de hoy
        fecha_hoy = datetime.now().strftime('%d/%m/%Y')
        print(f"📅 Filtrando noticias de hoy ({fecha_hoy})...")
        df_hoy = df[df['Fecha'] == fecha_hoy]
        
        # Convertir a formato de noticias
        noticias = []
        for _, row in df_hoy.iterrows():
            noticia = {
                'title': row['A'],  # Título en columna A
                'feed_name': 'whatsapp_group',
                'published': row['Fecha'],
                'link': ''  # No hay link disponible
            }
            noticias.append(noticia)
        
        print(f"✅ Google Sheets: {len(noticias)} noticias cargadas de hoy")
        return noticias
        
    except Exception as e:
        print(f"❌ Error al cargar noticias de Google Sheets: {str(e)}")
        print(f"Tipo de error: {type(e).__name__}")
        import traceback
        print("Detalles del error:")
        print(traceback.format_exc())
        return []

def procesar_feed(args):
    """Procesa un feed RSS individual"""
    name, url = args
    max_retries = 3
    retry_count = 0
    
    while retry_count < max_retries:
        try:
            rss_reader = RSSReader()
            rss_reader.add_feed(name, url)
            entries = rss_reader.get_feed_entries(name, limit=50)
            if entries:
                print(f"✅ {name}: {len(entries)} noticias cargadas")
                return name, entries
            else:
                print(f"⚠️ Advertencia: No se encontraron entradas en el feed {name}")
                return name, []
        except Exception as e:
            retry_count += 1
            if retry_count < max_retries:
                print(f"⚠️ Intento {retry_count} fallido para {name}: {str(e)}")
                time.sleep(2)
            else:
                print(f"❌ Error al procesar {name} después de {max_retries} intentos: {str(e)}")
                return name, []

def clasificar_noticia(noticia, temas):
    """Clasifica una noticia según los temas del diccionario"""
    titulo = noticia['title'].lower()
    temas_encontrados = []
    
    for tema_data in temas.values():
        # Por cada palabra clave del tema, verificar si está contenida en el título
        for palabra_clave in tema_data['palabras_clave']:
            if palabra_clave in titulo:
                tema_encontrado = {
                    'tema': tema_data['tema'],
                    'palabra_encontrada': palabra_clave
                }
                logging.info(f"Noticia: '{noticia['title']}' clasificada en tema: '{tema_data['tema']}' por palabra clave: '{palabra_clave}'")
                temas_encontrados.append(tema_encontrado)
                break
    
    return temas_encontrados

def generar_lista_medios(feeds):
    """Genera una lista formateada de los medios utilizados"""
    medios = {
        'infobae': 'Infobae',
        'clarin': 'Clarín',
        'lanacion_principal': 'La Nación',
        'lanacion_politica': 'La Nación - Política',
        'lanacion_economia': 'La Nación - Economía',
        'lanacion_deportes': 'La Nación - Deportes',
        'lanacion_sociedad': 'La Nación - Sociedad',
        'lanacion_mundo': 'La Nación - El Mundo',
        'lanacion_tecnologia': 'La Nación - Tecnología',
        'lanacion_opinion': 'La Nación - Opinión',
        'lanacion_lifestyle': 'La Nación - Lifestyle',
        'lagaceta_general': 'La Gaceta',
        'lagaceta_politica': 'La Gaceta - Política',
        'lagaceta_economia': 'La Gaceta - Economía',
        'lagaceta_deportes': 'La Gaceta - Deportes',
        'lagaceta_sociedad': 'La Gaceta - Sociedad',
        'lagaceta_mundo': 'La Gaceta - Mundo',
        'lagaceta_espectaculos': 'La Gaceta - Espectáculos',
        'lagaceta_opinion': 'La Gaceta - Opinión',
        'misiones_online': 'Misiones Online',
        'eldia': 'El Día',
        'rionegro': 'Río Negro',
        'diario_cuyo': 'Diario de Cuyo',
        'pagina12': 'Página/12',
        'cronista': 'El Cronista',
        'tn': 'TN',
        'ambito_home': 'Ámbito Financiero',
        'ambito_economia': 'Ámbito - Economía',
        'ambito_ultimas': 'Ámbito - Últimas Noticias',
        'ambito_politica': 'Ámbito - Política',
        'ambito_tecnologia': 'Ámbito - Tecnología',
        'lavoz_principal': 'La Voz',
        'lavoz_noticias': 'La Voz - Noticias',
        'lavoz_politica': 'La Voz - Política',
        'lavoz_negocios': 'La Voz - Negocios',
        'lavoz_ciudadanos': 'La Voz - Ciudadanos',
        'lavoz_deportes': 'La Voz - Deportes',
        'lavoz_cultura': 'La Voz - Cultura',
        'lavoz_internacionales': 'La Voz - Internacionales',
        'lavoz_espectaculos': 'La Voz - Espectáculos',
        'lavoz_opinion': 'La Voz - Opinión'
    }
    
    # Obtener medios únicos (sin secciones)
    medios_unicos = set()
    for feed_id in feeds.keys():
        # Obtener el nombre base del medio (antes del guion bajo)
        medio_base = feed_id.split('_')[0]
        if medio_base in medios:
            medios_unicos.add(medios[medio_base])
    
    return sorted(list(medios_unicos))

def guardar_en_word(noticias_por_tema, total_noticias, tiempo_total, feeds):
    """Guarda los resultados en un documento Word"""
    doc = Document()
    
    # Configurar el título
    titulo = doc.add_heading('Reporte de Noticias por Tema', 0)
    titulo.alignment = 1  # Centrado
    
    # Agregar información general
    doc.add_paragraph(f'Fecha del reporte: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph(f'Total de noticias clasificadas: {total_noticias}')
    doc.add_paragraph(f'Tiempo de procesamiento: {tiempo_total:.2f} segundos')
    
    # Agregar lista de medios consultados
    doc.add_heading('Medios consultados:', level=1)
    medios_list = generar_lista_medios(feeds)
    for medio in medios_list:
        doc.add_paragraph(f'• {medio}')
    
    # Agregar noticias por tema
    for tema, noticias in noticias_por_tema.items():
        if noticias:
            # Agregar encabezado del tema
            doc.add_heading(f'{tema} ({len(noticias)} noticias)', level=1)
            
            # Agregar cada noticia
            for noticia in noticias:
                p = doc.add_paragraph()
                p.add_run(f'• {noticia["title"]}').bold = True
                p.add_run(f'\n  {noticia["feed_name"]} - {noticia.get("published", "Fecha no disponible")}')
                p.add_run(f'\n  Palabra clave encontrada: {noticia["palabra_encontrada"]}')
                p.add_run(f'\n  {noticia["link"]}')
                doc.add_paragraph()  # Espacio entre noticias
    
    # Crear directorio si no existe
    os.makedirs('Reportes', exist_ok=True)
    
    # Guardar el documento
    fecha_actual = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_archivo = f'Reportes/reporte_noticias_{fecha_actual}.docx'
    doc.save(nombre_archivo)
    print(f"\n✅ Reporte guardado como: {nombre_archivo}")

def cargar_cuentas():
    """Carga la información de cuentas y sus temas asociados desde el Excel"""
    try:
        logging.info("📂 Intentando abrir Cuentas.xlsx...")
        
        # Verificar si el archivo existe
        if not os.path.exists('Cuentas.xlsx'):
            logging.error("❌ El archivo Cuentas.xlsx no existe en el directorio actual")
            return {}
            
        # Intentar leer el archivo
        df = pd.read_excel('Cuentas.xlsx', engine='openpyxl')
        logging.info("✅ Archivo abierto correctamente")
        
        if df.empty:
            logging.error("❌ El archivo está vacío")
            return {}
            
        if 'Empresa' not in df.columns:
            logging.error("❌ El archivo no tiene la columna 'Empresa'")
            return {}
            
        logging.info(f"Columnas encontradas: {df.columns.tolist()}")
        
        # Diccionario para almacenar la relación empresa -> temas
        cuentas = {}
        
        # Procesar cada fila (empresa)
        for idx, row in df.iterrows():
            try:
                empresa = row['Empresa']
                
                # Si no hay empresa, saltamos
                if pd.isna(empresa):
                    continue
                    
                empresa = str(empresa).strip()
                
                # Inicializar la empresa en el diccionario
                if empresa not in cuentas:
                    cuentas[empresa] = {
                        'nombre': empresa,
                        'temas': []
                    }
                
                # Procesar cada columna de temáticas
                for columna in df.columns:
                    if columna != 'Empresa' and pd.notna(row[columna]):
                        tema = str(row[columna]).strip()
                        if tema and tema not in cuentas[empresa]['temas']:
                            logging.info(f"Empresa: {empresa} - Tema encontrado: '{tema}'")
                            cuentas[empresa]['temas'].append(tema)
                            
            except Exception as row_error:
                logging.error(f"Error procesando fila {idx}: {str(row_error)}")
                continue
        
        return cuentas
        
    except Exception as e:
        logging.error(f"❌ Error al cargar el archivo de cuentas: {str(e)}")
        import traceback
        logging.error("Detalles del error:")
        logging.error(traceback.format_exc())
        return {}

def guardar_reporte_cuenta(cuenta, temas_cuenta, noticias_por_tema, tiempo_total, feeds):
    """Guarda un reporte Word para una cuenta específica"""
    try:
        # Crear directorio para reportes
        fecha_str = datetime.now().strftime("%Y%m%d")
        dir_reportes = os.path.join("Reportes", fecha_str)
        os.makedirs(dir_reportes, exist_ok=True)
        
        # Crear documento Word
        doc = Document()
        
        # Título
        doc.add_heading(f'Reporte de Noticias - {cuenta}', 0)
        
        # Información general
        doc.add_paragraph(f'Fecha del reporte: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        doc.add_paragraph(f'Temas monitoreados: {", ".join(temas_cuenta)}')
        doc.add_paragraph(f'Medios relevados: {", ".join(generar_lista_medios(feeds))}')
        doc.add_paragraph()
        
        # Agregar noticias por tema
        for tema in temas_cuenta:
            if tema in noticias_por_tema and noticias_por_tema[tema]:
                doc.add_heading(f'📌 {tema.upper()}', level=1)
                doc.add_paragraph('_' * 80)
                
                for noticia in noticias_por_tema[tema]:
                    p = doc.add_paragraph()
                    p.add_run(f'• {noticia["title"]}').bold = True
                    p.add_run(f'\n  {noticia["feed_name"]} - {noticia["published"]}')
                    if "palabra_encontrada" in noticia:
                        p.add_run(f'\n  Palabra clave encontrada: {noticia["palabra_encontrada"]}')
                    if "link" in noticia:
                        p.add_run(f'\n  {noticia["link"]}')
                    doc.add_paragraph()
        
        # Agregar contenido del Boletín Oficial si existe
        try:
            # Buscar el último reporte del BO
            dir_bo = "datos_boletin"
            if os.path.exists(dir_bo):
                # Encontrar la última carpeta de reportes
                carpetas_reportes = [d for d in os.listdir(dir_bo) if d.endswith('_reportes_cuenta')]
                if carpetas_reportes:
                    ultima_carpeta = sorted(carpetas_reportes)[-1]
                    ruta_reporte_bo = os.path.join(dir_bo, ultima_carpeta, f"reporte_detallado_{cuenta.replace(' ', '_').replace('/', '_')}.docx")
                    
                    if os.path.exists(ruta_reporte_bo):
                        # Agregar separador
                        doc.add_heading('=' * 50, level=1)
                        doc.add_heading('DOCUMENTOS DEL BOLETÍN OFICIAL', level=1)
                        doc.add_paragraph('_' * 80)
                        
                        # Cargar el documento del BO
                        doc_bo = Document(ruta_reporte_bo)
                        
                        # Copiar el contenido desde después del título principal
                        for element in list(doc_bo.element.body)[2:]:  # Saltar título y fecha
                            doc.element.body.append(element)
                            
                        doc.add_paragraph(f"\nInformación extraída del reporte: {ruta_reporte_bo}")
                        
        except Exception as e:
            doc.add_paragraph(f"\nNota: No se pudo incluir información del Boletín Oficial. Error: {str(e)}")
        
        # Guardar documento
        nombre_archivo = f"reporte_{cuenta.replace(' ', '_')}_{datetime.now().strftime('%H%M%S')}.docx"
        ruta_archivo = os.path.join(dir_reportes, nombre_archivo)
        doc.save(ruta_archivo)
        
        print(f"✅ Reporte para {cuenta} guardado como: {ruta_archivo}")
        return True
        
    except Exception as e:
        print(f"❌ Error al generar reporte para {cuenta}: {str(e)}")
        return False

def main():
    print("🔄 Iniciando lectura de feeds RSS...")
    
    # Crear carpeta principal de reportes si no existe
    os.makedirs('Reportes', exist_ok=True)
    
    # Cargar cuentas y sus temas
    print("\n📊 Cargando información de cuentas...")
    cuentas = cargar_cuentas()
    if not cuentas:
        print("❌ No se pudo cargar el archivo de cuentas. Saliendo...")
        return
    
    print(f"✅ Se cargaron {len(cuentas)} cuentas")
    print("\nCuentas y temas cargados:")
    for cuenta, data in cuentas.items():
        print(f"• {cuenta}: {len(data['temas'])} temas")
        for tema in data['temas']:
            print(f"  - {tema}")
    
    # Cargar temas del diccionario
    print("\n📚 Cargando diccionario de temas...")
    temas = cargar_diccionario()
    if not temas:
        print("❌ No se pudo cargar el diccionario de temas. Saliendo...")
        return
    
    print(f"✅ Se cargaron {len(temas)} temas del diccionario")
    print("\nTemas cargados:")
    for tema, data in temas.items():
        print(f"• {tema} ({len(data['palabras_clave'])} palabras clave)")
    
    # Configurar las fuentes RSS
    feeds = {
        'infobae': 'https://www.infobae.com/feeds/rss/',
        'clarin': 'https://www.clarin.com/rss/lo-ultimo/',
        'lanacion_principal': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/?outputType=xml',
        'lanacion_politica': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/politica/?outputType=xml',
        'lanacion_economia': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/economia/?outputType=xml',
        'lanacion_deportes': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/deportes/?outputType=xml',
        'lanacion_sociedad': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/sociedad/?outputType=xml',
        'lanacion_mundo': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/el-mundo/?outputType=xml',
        'lanacion_tecnologia': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/tecnologia/?outputType=xml',
        'lanacion_opinion': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/opinion/?outputType=xml',
        'lanacion_lifestyle': 'https://www.lanacion.com.ar/arc/outboundfeeds/rss/categoria/lifestyle/?outputType=xml',
        # La Política Online
        'lpo_ultimas': 'http://www.lapoliticaonline.com.ar/files/rss/ultimasnoticias.xml',
        'lpo_politica': 'http://www.lapoliticaonline.com.ar/files/rss/politica.xml',
        'lpo_economia': 'http://www.lapoliticaonline.com.ar/files/rss/economia.xml',
        # Feeds de medios regionales
        'misiones_online': 'https://misionesonline.net/feed/',
        'eldia': 'https://www.eldia.com/.rss',
        'rionegro': 'https://www.rionegro.com.ar/feed/',
        'diario_cuyo': 'https://www.diariodecuyo.com.ar/rss/rss.xml',
        # Feeds de La Gaceta
        'lagaceta_general': 'https://feeds.feedburner.com/LaGaceta-General',
        'lagaceta_politica': 'https://www.lagaceta.com.ar/rss/politica.xml',
        'lagaceta_economia': 'https://www.lagaceta.com.ar/rss/economia.xml',
        'lagaceta_deportes': 'https://www.lagaceta.com.ar/rss/deportes.xml',
        'lagaceta_sociedad': 'https://www.lagaceta.com.ar/rss/sociedad.xml',
        'lagaceta_mundo': 'https://www.lagaceta.com.ar/rss/mundo.xml',
        'lagaceta_espectaculos': 'https://www.lagaceta.com.ar/rss/espectaculos.xml',
        'lagaceta_opinion': 'https://www.lagaceta.com.ar/rss/opinion.xml',
        'pagina12': 'https://www.pagina12.com.ar/rss/portada',
        'cronista': 'https://www.cronista.com/files/rss/news.xml',
        'tn': 'https://tn.com.ar/rss',
        'mdz_politica': 'https://www.mdzol.com/rss/feed.html?r=1',
        'mdz_opinion': 'https://www.mdzol.com/rss/feed.html?r=4',
        'mdz_dinero': 'https://www.mdzol.com/rss/feed.html?r=5',
        'mdz_mundo': 'https://www.mdzol.com/rss/feed.html?r=6',
        'mdz_sociedad': 'https://www.mdzol.com/rss/feed.html?r=7',
        'mdz_deportes': 'https://www.mdzol.com/rss/feed.html?r=8',
        'mdz_sociales': 'https://www.mdzol.com/rss/feed.html?r=9',
        'mdz_show': 'https://www.mdzol.com/rss/feed.html?r=10',
        'mdz_espectaculos': 'https://www.mdzol.com/rss/feed.html?r=82',
        'mdz_policiales': 'https://www.mdzol.com/rss/feed.html?r=84',
        'mdz_energia': 'https://www.mdzol.com/rss/feed.html?r=90',
        # Feeds de Ámbito Financiero
        'ambito_home': 'https://www.ambito.com/rss/pages/home.xml',
        'ambito_economia': 'https://www.ambito.com/rss/pages/economia.xml',
        'ambito_ultimas': 'https://www.ambito.com/rss/pages/ultimas-noticias.xml',
        'ambito_politica': 'https://www.ambito.com/rss/pages/politica.xml',
        'ambito_tecnologia': 'https://www.ambito.com/rss/pages/tecnologia.xml',
        # Feeds de La Voz
        'lavoz_principal': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?origen=1',
        'lavoz_noticias': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?origen=2',
        'lavoz_politica': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=1',
        'lavoz_negocios': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=48',
        'lavoz_ciudadanos': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=17',
        'lavoz_deportes': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=18',
        'lavoz_cultura': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=16',
        'lavoz_internacionales': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=47',
        'lavoz_espectaculos': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=213',
        'lavoz_opinion': 'http://archivo.lavoz.com.ar/RSS/RSS.asp?categoria=214',
    }
    
    all_news = []
    start_time = time.time()
    
    # Cargar noticias de Google Sheets
    sheet_news = cargar_noticias_sheets()
    all_news.extend(sheet_news)
    
    # Procesar feeds en paralelo
    print("\n🔄 Procesando feeds en paralelo...")
    executor = concurrent.futures.ThreadPoolExecutor(max_workers=3)
    try:
        futures = [executor.submit(procesar_feed, (name, url)) for name, url in feeds.items()]
        
        for future in concurrent.futures.as_completed(futures):
            try:
                name, entries = future.result()
                if entries:
                    for entry in entries:
                        entry['feed_name'] = name
                        all_news.append(entry)
            except Exception as e:
                print(f"❌ Error inesperado: {str(e)}")
    finally:
        print("🔄 Cerrando threads...")
        executor.shutdown(wait=True)
        print("✅ Threads cerrados correctamente")
    
    tiempo_total = time.time() - start_time
    print(f"\n✨ Carga completada en {tiempo_total:.2f} segundos")
    
    # Clasificar noticias por tema
    print("\n📊 Clasificando noticias por tema...")
    noticias_por_tema = defaultdict(list)
    total_noticias_clasificadas = 0
    
    for noticia in all_news:
        temas_noticia = clasificar_noticia(noticia, temas)
        for tema_info in temas_noticia:
            tema = tema_info['tema']
            noticia['palabra_encontrada'] = tema_info['palabra_encontrada']
            noticias_por_tema[tema].append(noticia)
            total_noticias_clasificadas += 1
    
    # Mostrar resultados en consola
    print("\n=== NOTICIAS POR TEMA ===")
    for tema, noticias in noticias_por_tema.items():
        if noticias:
            print(f"\n📌 {tema.upper()} ({len(noticias)} noticias)")
            print("-" * 80)
            
            for noticia in noticias:
                fecha = noticia.get('published', 'Fecha no disponible')
                print(f"\n• {noticia['title']}")
                print(f"  {noticia['feed_name']} - {fecha}")
                print(f"  Palabra clave encontrada: {noticia['palabra_encontrada']}")
                print(f"  {noticia['link']}")
            print("=" * 80)
    
    print(f"\n📊 Total de noticias clasificadas: {total_noticias_clasificadas}")
    
    # Generar reporte general
    print("\n📝 Generando reporte general...")
    guardar_en_word(noticias_por_tema, total_noticias_clasificadas, tiempo_total, feeds)
    
    # Generar reportes por cuenta
    print("\n📝 Generando reportes por cuenta...")
    reportes_generados = 0
    for cuenta, data in cuentas.items():
        print(f"\n📌 Procesando cuenta: {cuenta}")
        guardar_reporte_cuenta(cuenta, data['temas'], noticias_por_tema, tiempo_total, feeds)
        reportes_generados += 1
    
    print(f"\n✨ Proceso completado. Se generaron reportes para {reportes_generados} cuentas con noticias relevantes.")

if __name__ == "__main__":
    main() 